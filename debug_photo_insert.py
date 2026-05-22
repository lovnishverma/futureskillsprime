from docx import Document
from docx.shared import Inches
import os

try:
    doc = Document("docxtemplates/GOT_Nomination_Form.docx")
    
    # We will find the cell with "{Photo}" and replace it with a local dummy image
    
    # Generate a local dummy image
    from PIL import Image
    dummy = Image.new('RGB', (200, 250), color='green')
    dummy_path = "dummy_photo.jpg"
    dummy.save(dummy_path)
    
    found = False
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                if "{Photo}" in cell.text:
                    found = True
                    for para in cell.paragraphs:
                        para.clear()
                    run = cell.paragraphs[0].add_run()
                    run.add_picture(dummy_path, width=Inches(1.2))
                    print("Added picture to cell.")
    
    doc.save("test_photo_out.docx")
    print("Found:", found)
    
except Exception as e:
    print("Error:", e)
