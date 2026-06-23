"""
NIELIT Nomination Form - Full Flask App
Features:
  - Online nomination form with photo + signature upload
  - Generates filled DOCX + PDF per submission
  - Participant can download their own PDF
  - Admin panel: view all submissions, download CSV, download individual or all-in-one PDF
"""

from reportlab.platypus import (Image as RLImage, Paragraph, SimpleDocTemplate,
                                Spacer, Table, TableStyle)
from reportlab.lib.units import cm, mm, inch
from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
from reportlab.lib.pagesizes import A4
from reportlab.lib import colors
from pypdf import PdfReader, PdfWriter
from PIL import Image
from flask import (Flask, abort, flash, g, redirect, render_template,
                   request, send_file, session, url_for)
from docx.shared import Inches, Pt
from docx import Document
from docxcompose.composer import Composer
from pathlib import Path
from io import BytesIO
from datetime import datetime
import zipfile
import uuid
import tempfile
from bson.objectid import ObjectId
import cloudinary.uploader
import cloudinary
from pymongo import MongoClient
import re
import os
import logging
import json
import io
from dotenv import load_dotenv
load_dotenv()


# ── App setup ────────────────────────────────────────────────────────────────
BASE_DIR = Path(__file__).resolve().parent
app = Flask(__name__)
app.secret_key = os.environ.get("FLASK_SECRET_KEY", "nielit_dev_secret_2026")
app.config["MAX_CONTENT_LENGTH"] = 10 * 1024 * 1024  # 10 MB
app.config["UPLOAD_FOLDER"] = BASE_DIR / "static" / "uploads"
app.config["MONGO_URI"] = os.environ.get("MONGO_URI")
app.config["ADMIN_PASSWORD"] = os.environ.get("ADMIN_PASSWORD")

cloudinary.config(
    cloud_name="dyq802zwa",
    api_key=os.environ.get("CLOUDINARY_API_KEY", ""),
    api_secret=os.environ.get("CLOUDINARY_API_SECRET", "")
)

DOCX_TEMPLATE = BASE_DIR / "docxtemplates" / "GOT_Nomination_Form.docx"
ALLOWED_IMG = {"png", "jpg", "jpeg", "gif", "webp"}

logging.basicConfig(level=logging.INFO)


# ── Database ──────────────────────────────────────────────────────────────────
mongo_client = MongoClient(app.config["MONGO_URI"])
db_client = mongo_client["nielit_db"]
nominations_col = db_client["nominations"]


def get_db():
    return nominations_col


def init_db():
    nominations_col.create_index("token", unique=True)

# Call it immediately when the module is imported by Gunicorn
init_db()


# ── Helpers ───────────────────────────────────────────────────────────────────




import urllib.request


def fetch_image_as_jpeg(url, target_size=None):
    import urllib.request
    import tempfile
    from io import BytesIO
    from PIL import Image, ImageOps
    req = urllib.request.urlopen(url)
    img = Image.open(BytesIO(req.read()))
    if img.mode in ("RGBA", "P"):
        img = img.convert("RGB")
        
    if target_size:
        # Replaced ImageOps.fit with ImageOps.contain
        img = ImageOps.contain(img, target_size, Image.Resampling.LANCZOS if hasattr(Image, 'Resampling') else Image.ANTIALIAS)
        
    fd, path = tempfile.mkstemp(suffix=".jpg")
    import os
    os.close(fd)
    img.save(path, format="JPEG")
    return path

def upload_to_cloudinary(file_obj, folder_name):
    """Upload file directly to Cloudinary and return secure URL."""
    if not file_obj or file_obj.filename == '':
        return None
    try:
        res = cloudinary.uploader.upload(file_obj, folder=folder_name, resource_type="image")
        return res.get("secure_url")
    except Exception as e:
        logging.error(f"Cloudinary upload failed: {e}")
        return None

def fmt_date(val):
    """Convert YYYY-MM-DD → DD-MM-YYYY."""
    if val and re.match(r"^\d{4}-\d{2}-\d{2}$", val):
        try:
            return datetime.strptime(val, "%Y-%m-%d").strftime("%d-%m-%Y")
        except Exception:
            pass
    return val or ""

def get_ordinal(n):
    return str(n) + ('th' if 11 <= n % 100 <= 13 else {1: 'st', 2: 'nd', 3: 'rd'}.get(n % 10, 'th'))

def fmt_course_dates(start_str, end_str):
    if not start_str or not end_str:
        return ""
    try:
        s = datetime.strptime(start_str, "%Y-%m-%d")
        e = datetime.strptime(end_str, "%Y-%m-%d")
        s_fmt = f"{get_ordinal(s.day)} {s.strftime('%B')}, {s.year}"
        e_fmt = f"{get_ordinal(e.day)} {e.strftime('%B')}, {e.year}"
        return f"{s_fmt} - {e_fmt}"
    except Exception:
        return ""

def row_to_form_data(row):
    """Map DB row to the placeholder dict used by DOCX/PDF generation."""
    d = dict(row)
    name_full = f"{d.get('title','')} {d.get('name','')}".strip()
    
    track = (d.get("track") or "").upper()
    level = (d.get("level") or "").capitalize()
    course_key = f"{track}_{level}"
    
    c_start = d.get("course_start_date")
    c_end = d.get("course_end_date")
    
    if not c_start or not c_end:
        # Fallback for old nominations that didn't save dates directly
        config_col = db_client["config"]
        course_dates_doc = config_col.find_one({"_id": "course_dates"}) or {}
        batches = course_dates_doc.get(course_key, [])
        if isinstance(batches, dict):
            batches = [batches] if batches.get("start") else []
        if batches:
            c_start = c_start or batches[0].get("start")
            c_end = c_end or batches[0].get("end")

    formatted_dates = fmt_course_dates(c_start, c_end) if c_start and c_end else ""
    
    return {
        "Title": d.get("title", ""),
        "Name": d.get("name", ""),
        "Full_Name": name_full,
        "Course_Name": _course_name(d.get("track", ""), d.get("level", "")),
        "Course_Start_Date": formatted_dates,
        "Date_of_Training": formatted_dates,
        "Native_State": d.get("native_state", ""),
        "District": d.get("district", ""),
        "Status": d.get("status", ""),
        "Beneficiary_Category": d.get("beneficiary_category", ""),
        "Level": d.get("level", ""),
        "Applicant_Name": name_full,
        "Gov_ID_Number": d.get("aadhar", ""),
        "Organization_Academic_Institute": d.get("organisation", ""),
        "Role": d.get("designation", ""),
        "Highest_Qualification": f"{d.get('highest_qual', '')}",
        "Previous_FSP_Program": d.get("prev_fsp", ""),
        "Previous_FSP_Details_1": d.get("prev_fsp_details", ""),
        "Previous_FSP_Details_2": "",
        "Technology": _technology(d.get("track", "")),
        "Resource_Centre_Name": d.get("resource_centre", "NIELIT Chandigarh"),
        "Date_of_Training": formatted_dates,
        "Designation": d.get("designation", ""),
        "Organisation": d.get("organisation", ""),
        "DOB": fmt_date(d.get("dob", "")),
        "Gender": d.get("gender", ""),
        "Aadhar": d.get("aadhar", ""),
        "Contact_Number": d.get("contact", ""),
        "Email": d.get("email", ""),
        "Organisation_Department": _merge(d.get("organisation", ""), d.get("department", "")),
        "Contact_Number_Email": _merge(d.get("contact", ""), d.get("email", "")),
        "Institute_Details": _merge(_merge(d.get("institute_address", ""), d.get("institute_contact", "")), d.get("institute_email", "")),
        "photo_url": d.get("photo_url", ""),
        "sign_url": d.get("sign_url", ""),
        "Edu1_Year": d.get("edu1_year", ""), "Edu1_Degree": d.get("edu1_degree", ""), "Edu1_University": d.get("edu1_university", ""),
        "Edu2_Year": d.get("edu2_year", ""), "Edu2_Degree": d.get("edu2_degree", ""), "Edu2_University": d.get("edu2_university", ""),
        "Edu3_Year": d.get("edu3_year", ""), "Edu3_Degree": d.get("edu3_degree", ""), "Edu3_University": d.get("edu3_university", ""),
        "Exp1_Year": d.get("exp1_year", ""), "Exp1_Area_of_Expertise": d.get("exp1_area", ""), "Exp1_Centre": d.get("exp1_centre", ""),
        "Exp2_Year": d.get("exp2_year", ""), "Exp2_Area_of_Expertise": d.get("exp2_area", ""), "Exp2_Centre": d.get("exp2_centre", ""),
        "Exp3_Year": d.get("exp3_year", ""), "Exp3_Area_of_Expertise": d.get("exp3_area", ""), "Exp3_Centre": d.get("exp3_centre", ""),
        "Today_Date": datetime.today().strftime("%d-%m-%Y"),
        "photo_url": d.get("photo_url"),
        "sign_url": d.get("sign_url"),
    }


def _course_name(track, level):
    track = (track or "").upper()
    level = (level or "").capitalize()
    names = {
        ("ARVR", "Basic"):    "GOT – AR & VR (Basic)",
        ("ARVR", "Advanced"): "GOT – AR & VR (Advanced)",
        ("ARVR", "Bootcamp"): "Bootcamp – AR & VR",
        ("BDDS", "Basic"):    "GOT – Big Data & Data Science (Basic)",
        ("BDDS", "Advanced"): "GOT – Big Data & Data Science (Advanced)",
        ("BDDS", "Bootcamp"): "Bootcamp – Big Data & Data Science",
    }
    return names.get((track, level), f"{track} {level}")


def _technology(track):
    t = (track or "").upper()
    return {"ARVR": "Augmented & Virtual Reality", "BDDS": "Big Data & Data Science"}.get(t, track)


def _merge(a, b, sep=" / "):
    parts = [x for x in [a, b] if x]
    return sep.join(parts) if parts else ""


# ── DOCX generation ───────────────────────────────────────────────────────────
def _replace_in_para(para, replacements):
    full = para.text
    if "{" not in full and "<<" not in full:
        return
    if para.runs:
        for run in para.runs:
            for k, v in replacements.items():
                run.text = run.text.replace(
                    f"{{{k}}}", v).replace(f"<<{k}>>", v)
        joined = "".join(r.text for r in para.runs)
        if "{" in joined or "<<" in joined:
            updated = joined
            for k, v in replacements.items():
                updated = updated.replace(f"{{{k}}}", v).replace(f"<<{k}>>", v)
            if updated != joined:
                para.runs[0].text = updated
                for r in para.runs[1:]:
                    r.text = ""


def generate_docx(form_data: dict) -> BytesIO:
    """Fill the DOCX template and return as BytesIO. Inserts photo into Photo cell."""
    template_path = Path("docxtemplates/Bootcamp_Nomination_Form.docx") if form_data.get("Level") == "Bootcamp" else DOCX_TEMPLATE
    doc = Document(str(template_path))
    replacements = {k: (str(v) if v else "") for k, v in form_data.items()
                    if k not in ("photo_url", "sign_url")}

    level_val = (form_data.get("Level") or "").capitalize()
    if level_val in ("Advanced", "Basic"):
        prog_title = f"Government Officials Training ({level_val}) Programme"
    else:
        prog_title = "Government Officials Training Programme"

    for para in doc.paragraphs:
        if "Government Officials Training Programme" in para.text:
            for run in para.runs:
                run.text = run.text.replace("Government Officials Training Programme", prog_title)

        cell_text = para.text.strip()
        if "{Photo}" in cell_text or "<<Photo>>" in cell_text or cell_text == "Photo":
            para.text = "Photo"
            continue
        elif "{Signature}" in cell_text or "<<Signature>>" in cell_text or cell_text == "Signature":
            para.text = "[Please Sign Here]"
            continue
        
        if "Any of the government issued ID" in para.text:
            para.style = 'Normal'
            para.text = ""
            r0 = para.add_run("• ")
            r1 = para.add_run("1")
            r1.font.superscript = True
            r2 = para.add_run(" Any of the government issued ID: Aadhaar card")
            r2.font.italic = True
            continue
        elif "Participants are not allowed to enroll" in para.text:
            para.style = 'Normal'
            para.text = ""
            r0 = para.add_run("• ")
            r1 = para.add_run("2")
            r1.font.superscript = True
            r2 = para.add_run(" Participants are not allowed to enroll in the same program for more than once")
            r2.font.italic = True
            continue
        
        _replace_in_para(para, replacements)

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                # Check if this cell is the "Photo" placeholder cell
                cell_text = cell.text.strip()
                if "{Photo}" in cell_text or "<<Photo>>" in cell_text or cell_text == "Photo":
                    cell.paragraphs[0].text = "Photo"
                    continue
                elif "{Signature}" in cell_text or "<<Signature>>" in cell_text or cell_text == "Signature":
                    cell.paragraphs[0].text = "[Please Sign Here]"
                    continue
                for para in cell.paragraphs:
                    if "Any of the government issued ID" in para.text:
                        para.style = 'Normal'
                        para.text = ""
                        r0 = para.add_run("• ")
                        r1 = para.add_run("1")
                        r1.font.superscript = True
                        r2 = para.add_run(" Any of the government issued ID: Aadhaar card")
                        r2.font.italic = True
                        continue
                    elif "Participants are not allowed to enroll" in para.text:
                        para.style = 'Normal'
                        para.text = ""
                        r0 = para.add_run("• ")
                        r1 = para.add_run("2")
                        r1.font.superscript = True
                        r2 = para.add_run(" Participants are not allowed to enroll in the same program for more than once")
                        r2.font.italic = True
                        continue
                        
                    _replace_in_para(para, replacements)

    for section in doc.sections:
        for para in section.header.paragraphs:
            _replace_in_para(para, replacements)
        for para in section.footer.paragraphs:
            _replace_in_para(para, replacements)

    buf = BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf


# ── PDF generation via ReportLab ──────────────────────────────────────────────
def generate_pdf(form_data: dict) -> BytesIO:
    """Generate a simple A4 PDF nomination form matching the DOCX template."""
    buf = BytesIO()
    
    if form_data.get("Level") == "Bootcamp":
        doc = SimpleDocTemplate(
            buf, pagesize=A4,
            topMargin=0.8*cm, bottomMargin=0.8*cm,
            leftMargin=1.2*cm, rightMargin=1.2*cm,
            title="Nomination Form",
            author="FutureSkills PRIME"
        )
        elements = []
        
        # Styles
        styles = getSampleStyleSheet()
        title_st = ParagraphStyle('title', fontName='Helvetica-Bold', fontSize=12, alignment=1, spaceAfter=4)
        norm_st = ParagraphStyle('norm', fontName='Helvetica', fontSize=9)
        bold_st = ParagraphStyle('bold', fontName='Helvetica-Bold', fontSize=9)
        
        usable_w = A4[0] - 2.4*cm
        
        # Header with Logos
        try:
            with open("static/assets/img/NIELIT-Logo-hd.png", "rb") as f:
                n_data = BytesIO(f.read())
            nielit_img = RLImage(n_data, width=1.6*inch, height=0.85*inch)
        except Exception:
            nielit_img = ""
            
        try:
            fs_img = RLImage("static/assets/img/FUTURESKILLS_LOGO.png", width=1.5*inch, height=0.7*inch)
        except Exception:
            fs_img = ""
            
        header_text = """
        <font color="#000080" size="14"><b>National Institute of Electronics and<br/>Information Technology Chandigarh</b></font><br/>
        <font color="#000080" size="7">An autonomous scientific society under administrative control of<br/>
        Ministry of Electronics and Information Technology (MeitY), Government of India</font>
        """
        header_p = Paragraph(header_text, ParagraphStyle('h', alignment=1, leading=10))
        
        header_t = Table([[nielit_img, header_p, fs_img]], colWidths=[1.5*inch, usable_w-3.0*inch, 1.5*inch])
        header_t.setStyle(TableStyle([
            ('VALIGN', (0,0), (-1,-1), 'MIDDLE'),
            ('ALIGN', (0,0), (-1,-1), 'CENTER'),
        ]))
        elements.append(header_t)
        
        fsp_text = "<b>FutureSkills PRIME (Programme for Re-skilling/Up-skilling of ITManpower for Employability)</b>"
        elements.append(Paragraph(fsp_text, ParagraphStyle('fsp', fontName='Helvetica-Bold', fontSize=10, alignment=1)))
        elements.append(Spacer(1, 15))
        
        elements.append(Paragraph("<b>Application Form for Bootcamp</b>", ParagraphStyle('afb', fontName='Helvetica-Bold', fontSize=14, alignment=1)))
        elements.append(Spacer(1, 15))
        
        elements.append(Paragraph("<b>Date:</b>", ParagraphStyle('rdate', fontName='Helvetica-Bold', fontSize=10, alignment=2, rightIndent=100)))
        elements.append(Spacer(1, 5))
        
        data = [
            [Paragraph("<b>Resource Centre Name</b>", norm_st), "", ":", Paragraph(form_data.get("Resource_Centre_Name", ""), norm_st), ""],
            [Paragraph("<b>Technology</b>", norm_st), "", ":", Paragraph(form_data.get("Technology", ""), norm_st), Paragraph("<b>Role:</b> Co-Lead", norm_st)],
            ["1", Paragraph("Course Name", norm_st), ":", Paragraph(form_data.get("Course_Name", ""), norm_st), ""],
            ["2", Paragraph("Course Start Date", norm_st), ":", Paragraph(form_data.get("Course_Start_Date", ""), norm_st), ""],
            ["3", Paragraph("Applicant Name (as per Gov ID)", norm_st), ":", Paragraph(form_data.get("Applicant_Name", ""), norm_st), ""],
            ["4", Paragraph("Date of Birth", norm_st), ":", Paragraph(form_data.get("DOB", ""), norm_st), ""],
            ["5", Paragraph("Gender", norm_st), ":", Paragraph(form_data.get("Gender", ""), norm_st), ""],
            ["6", Paragraph("Mobile Number", norm_st), ":", Paragraph(form_data.get("Contact_Number", ""), norm_st), ""],
            ["7", Paragraph("Email ID (official ID preferred)", norm_st), ":", Paragraph(form_data.get("Email", ""), norm_st), ""],
            ["8", Paragraph("Native State", norm_st), ":", Paragraph(form_data.get("Native_State", ""), norm_st), ""],
            ["9", Paragraph("District", norm_st), ":", Paragraph(form_data.get("District", ""), norm_st), ""],
            ["10", Paragraph("Government-issued ID Number<sup>1</sup><br/><i>(Aadhar copy enclosed)</i>", norm_st), ":", Paragraph(form_data.get("Gov_ID_Number", ""), norm_st), ""],
            ["11", Paragraph("Organization/Academic Institute (if applicable)", norm_st), ":", Paragraph(form_data.get("Organization_Academic_Institute", ""), norm_st), ""],
            ["12", Paragraph("Highest Qualification (with Degree & Branch)", norm_st), ":", Paragraph(form_data.get("Highest_Qualification", ""), norm_st), ""],
            ["13", Paragraph("Status (Pursuing/ Passed out)", norm_st), ":", Paragraph(form_data.get("Status", ""), norm_st), ""],
            ["14", Paragraph("Beneficiary Category (Tick as applicable)", norm_st), ":", Paragraph(form_data.get("Beneficiary_Category", ""), norm_st), ""],
            ["15", Paragraph("Involved in previous FSP Program<sup>2</sup>", norm_st), ":", Paragraph(form_data.get("Previous_FSP_Program", ""), norm_st), ""],
            ["16", Paragraph("If previous answer is, yes provide details<br/>(Program Name / Conducting Institute / Date)", norm_st), ":", Paragraph(f"1. {form_data.get('Previous_FSP_Details_1', '')}<br/>2. {form_data.get('Previous_FSP_Details_2', '')}", norm_st), ""],
        ]
        
        t_style_cmds = [
            ('GRID', (0,2), (-1,-1), 0.5, colors.black),
            ('LINEABOVE', (0,0), (-1,0), 0.5, colors.black),
            ('LINEBELOW', (0,0), (-1,0), 0.5, colors.black),
            ('LINEBELOW', (0,1), (-1,1), 0.5, colors.black),
            ('VALIGN', (0,0), (-1,-1), 'MIDDLE'),
            ('PADDING', (0,0), (-1,-1), 4),
            ('SPAN', (0, 0), (1, 0)), 
            ('SPAN', (3, 0), (4, 0)), 
            ('SPAN', (0, 1), (1, 1)), 
        ]
        
        for r in range(2, len(data)):
            t_style_cmds.append(('SPAN', (3, r), (4, r)))
            
        t_style = TableStyle(t_style_cmds)
        
        t = Table(data, colWidths=[usable_w*0.05, usable_w*0.35, usable_w*0.05, usable_w*0.35, usable_w*0.20])
        t.setStyle(t_style)
        elements.append(t)
        elements.append(Spacer(1, 10))
        
        elements.append(Paragraph("I hereby declare that all the information provided above is true.", norm_st))
        
        elements.append(Spacer(1, 20))
        elements.append(Paragraph("<b>Applicant Signature</b>", ParagraphStyle('r', fontName='Helvetica', fontSize=9, alignment=2)))
        
        elements.append(Spacer(1, 5))
        long_line = Table([[""]], colWidths=[usable_w], rowHeights=[1])
        long_line.setStyle(TableStyle([('LINEABOVE', (0,0), (-1,-1), 0.5, colors.black)]))
        elements.append(long_line)
        
        elements.append(Spacer(1, 5))
        elements.append(Paragraph("<b>(For office purpose)</b><br/><br/>The above submitted information has been verified and recommended.", norm_st))
        
        elements.append(Spacer(1, 20))
        elements.append(Paragraph("(Signature)<br/><b>Course Co-ordinator</b>", ParagraphStyle('r2', fontName='Helvetica', fontSize=9, alignment=2)))
        
        elements.append(Spacer(1, 5))
        short_line = Table([[""]], colWidths=[150], rowHeights=[1], hAlign='LEFT')
        short_line.setStyle(TableStyle([('LINEABOVE', (0,0), (-1,-1), 0.5, colors.black)]))
        elements.append(short_line)
        
        elements.append(Spacer(1, 2))
        bullet_text = "&bull; &nbsp; <sup>1</sup> <i>Any of the government issued ID: Aadhaar card</i><br/>" \
                      "&bull; &nbsp; <sup>2</sup> <i>Participants are not allowed to enroll in the same program for more than once</i>"
        elements.append(Paragraph(bullet_text, ParagraphStyle('small', fontName='Helvetica', fontSize=8)))
        
        def draw_page_border(canvas, document):
            canvas.saveState()
            canvas.setStrokeColor(colors.black)
            canvas.setLineWidth(3)
            canvas.rect(0.6*cm, 0.6*cm, A4[0]-1.2*cm, A4[1]-1.2*cm)
            canvas.setLineWidth(0.5)
            canvas.rect(0.65*cm, 0.65*cm, A4[0]-1.3*cm, A4[1]-1.3*cm)
            canvas.restoreState()
            
        doc.build(elements, onFirstPage=draw_page_border, onLaterPages=draw_page_border)
        buf.seek(0)
        return buf
        
    # --- END BOOTCAMP BRANCH ---


    doc = SimpleDocTemplate(
        buf, pagesize=A4,
        topMargin=0.5*cm, bottomMargin=0.5*cm,
        leftMargin=1.2*cm, rightMargin=1.2*cm,
        title="Nomination Form",
        author="FutureSkills PRIME"
    )
    W, H = A4
    usable_w = W - 3.6*cm

    styles = getSampleStyleSheet()
    
    title_st = ParagraphStyle('title', fontName='Helvetica-Bold', fontSize=12, alignment=1, spaceAfter=4)
    sub_st = ParagraphStyle('sub', fontName='Helvetica-Bold', fontSize=10, alignment=1, spaceAfter=8)
    norm_st = ParagraphStyle('norm', fontName='Helvetica', fontSize=9)
    bold_st = ParagraphStyle('bold', fontName='Helvetica-Bold', fontSize=9)
    right_st = ParagraphStyle('right', fontName='Helvetica', fontSize=9, alignment=2)
    center_st = ParagraphStyle('center_st', fontName='Helvetica', fontSize=10, alignment=1)
    
    story = []
    
    try:
        with open("static/assets/img/NIELIT-Logo-hd.png", "rb") as f:
            n_data = BytesIO(f.read())
        n_img = RLImage(n_data, width=1.5*inch, height=0.8*inch)
        img_t = Table([[n_img]], colWidths=[usable_w])
        img_t.setStyle(TableStyle([('ALIGN', (0,0), (-1,-1), 'CENTER')]))
        story.append(img_t)
        story.append(Spacer(1, 10))
    except Exception:
        pass
        
    story.append(Paragraph("Nomination Form", title_st))
    
    level_val = (form_data.get("Level") or "").capitalize()
    if level_val in ("Advanced", "Basic"):
        prog_title = f"Government Officials Training ({level_val}) Programme"
    else:
        prog_title = "Government Officials Training Programme"
        
    story.append(Paragraph(prog_title, sub_st))
    
    # Table 0: Training Programme Details
    story.append(Paragraph("Training Programme Details", bold_st))
    t0_data = [
        [Paragraph("Name of Training Programme", norm_st), Paragraph(form_data.get("Course_Name", ""), norm_st)],
        [Paragraph("Name of Technology", norm_st), Paragraph(form_data.get("Technology", ""), norm_st)],
        [Paragraph("Resource Centre Name", norm_st), Paragraph(form_data.get("Resource_Centre_Name", ""), norm_st)],
        [Paragraph("Date of Training", norm_st), Paragraph(form_data.get("Date_of_Training", ""), norm_st)]
    ]
    t0 = Table(t0_data, colWidths=[usable_w*0.4, usable_w*0.6])
    t0.setStyle(TableStyle([
        ('BOX', (0,0), (-1,-1), 0.5, colors.black),
        ('INNERGRID', (0,0), (-1,-1), 0.5, colors.black),
        ('VALIGN', (0,0), (-1,-1), 'MIDDLE'),
        ('PADDING', (0,0), (-1,-1), 2)
    ]))
    story.append(t0)
    story.append(Spacer(1, 4))
    
    # Table 1: Personal Information
    story.append(Paragraph("Personal Information", bold_st))
    t1_data = [
        [Paragraph("NAME<br/>Prof./Dr./Mr./Ms.", norm_st), Paragraph(f"{form_data.get('Title','')} {form_data.get('Name','')}", norm_st), "", ""],
        [Paragraph("DESIGNATION:", norm_st), Paragraph(form_data.get("Designation", ""), norm_st), Paragraph("ORGANISATION:", norm_st), Paragraph(form_data.get("Organisation", ""), norm_st)],
        [Paragraph("DATE OF BIRTH :", norm_st), Paragraph(form_data.get("DOB", ""), norm_st), Paragraph("GENDER(M/F)", norm_st), Paragraph(form_data.get("Gender", ""), norm_st)],
        [Paragraph("AADHAR No.", norm_st), Paragraph(form_data.get("Aadhar", ""), norm_st), "", ""],
        [Paragraph("CONTACT NUMBER &<br/>E- MAIL", norm_st), Paragraph(f"{form_data.get('Contact_Number','')} / {form_data.get('Email','')}", norm_st), "", ""],
        [Paragraph("NAME OF THE<br/>ORGANISATION/<br/>DEPARTMENT", norm_st), Paragraph(form_data.get("Organisation_Department", ""), norm_st), "", ""],
        [Paragraph("COMPLETE ADDRESS /<br/>CONTACT NUMBERS /<br/>E- MAIL OF THE INSTITUTE", norm_st), Paragraph(form_data.get('Institute_Details', ''), norm_st), "", ""]
    ]
    t1 = Table(t1_data, colWidths=[usable_w*0.25, usable_w*0.25, usable_w*0.25, usable_w*0.25])
    t1.setStyle(TableStyle([
        ('BOX', (0,0), (-1,-1), 0.5, colors.black),
        ('INNERGRID', (0,0), (-1,-1), 0.5, colors.black),
        ('SPAN', (1,0), (3,0)),
        ('SPAN', (1,3), (3,3)),
        ('SPAN', (1,4), (3,4)),
        ('SPAN', (1,5), (3,5)),
        ('SPAN', (1,6), (3,6)),
        ('VALIGN', (0,0), (-1,-1), 'MIDDLE'),
        ('PADDING', (0,0), (-1,-1), 2)
    ]))
    story.append(t1)
    story.append(Spacer(1, 4))
    
    # Table 2: Educational Qualifications
    story.append(Paragraph("Educational / Professional Qualifications", bold_st))
    t2_data = [
        [Paragraph("EDUCATIONAL / PROFESSIONAL QUALIFICATIONS (GRADUATION ONWARDS)", bold_st), "", "", ""],
        [Paragraph("SL. No.", bold_st), Paragraph("YEAR", bold_st), Paragraph("DEGREE", bold_st), Paragraph("UNIVERSITY/INSTITUTE", bold_st)],
        ["1", Paragraph(form_data.get("Edu1_Year",""), norm_st), Paragraph(form_data.get("Edu1_Degree",""), norm_st), Paragraph(form_data.get("Edu1_University",""), norm_st)],
        ["2", Paragraph(form_data.get("Edu2_Year",""), norm_st), Paragraph(form_data.get("Edu2_Degree",""), norm_st), Paragraph(form_data.get("Edu2_University",""), norm_st)],
        ["3", Paragraph(form_data.get("Edu3_Year",""), norm_st), Paragraph(form_data.get("Edu3_Degree",""), norm_st), Paragraph(form_data.get("Edu3_University",""), norm_st)]
    ]
    t2 = Table(t2_data, colWidths=[usable_w*0.1, usable_w*0.2, usable_w*0.3, usable_w*0.4])
    t2.setStyle(TableStyle([
        ('BOX', (0,0), (-1,-1), 0.5, colors.black),
        ('INNERGRID', (0,0), (-1,-1), 0.5, colors.black),
        ('SPAN', (0,0), (3,0)),
        ('VALIGN', (0,0), (-1,-1), 'MIDDLE'),
        ('ALIGN', (0,0), (3,1), 'CENTER'),
        ('PADDING', (0,0), (-1,-1), 2)
    ]))
    story.append(t2)
    story.append(Spacer(1, 4))
    
    # Table 3: Experience
    t3_data = [
        [Paragraph("RESEARCH / TECHNICAL EXPERIENCE", bold_st), "", "", ""],
        [Paragraph("SL.NO.", bold_st), Paragraph("YEAR", bold_st), Paragraph("AREA OF EXPERTISE", bold_st), Paragraph("CENTRE", bold_st)],
        ["1", Paragraph(form_data.get("Exp1_Year",""), norm_st), Paragraph(form_data.get("Exp1_Area_of_Expertise",""), norm_st), Paragraph(form_data.get("Exp1_Centre",""), norm_st)],
        ["2", Paragraph(form_data.get("Exp2_Year",""), norm_st), Paragraph(form_data.get("Exp2_Area_of_Expertise",""), norm_st), Paragraph(form_data.get("Exp2_Centre",""), norm_st)],
        ["3", Paragraph(form_data.get("Exp3_Year",""), norm_st), Paragraph(form_data.get("Exp3_Area_of_Expertise",""), norm_st), Paragraph(form_data.get("Exp3_Centre",""), norm_st)]
    ]
    t3 = Table(t3_data, colWidths=[usable_w*0.1, usable_w*0.2, usable_w*0.4, usable_w*0.3])
    t3.setStyle(TableStyle([
        ('BOX', (0,0), (-1,-1), 0.5, colors.black),
        ('INNERGRID', (0,0), (-1,-1), 0.5, colors.black),
        ('SPAN', (0,0), (3,0)),
        ('VALIGN', (0,0), (-1,-1), 'MIDDLE'),
        ('ALIGN', (0,0), (3,1), 'CENTER'),
        ('PADDING', (0,0), (-1,-1), 2)
    ]))
    story.append(t3)
    story.append(Spacer(1, 5))
    
    # Table 4: Signature / Photo block
    photo_elem = Table([["Photo"]], colWidths=[1.1*inch], rowHeights=[1.1*inch], style=[('BOX', (0,0), (-1,-1), 1, colors.black), ('ALIGN', (0,0), (-1,-1), 'CENTER'), ('VALIGN', (0,0), (-1,-1), 'MIDDLE')])
            
    sign_text = (
        "<br/><br/><br/>"
        "<i>Signature of the Official</i><br/><br/>"
        "<b><i>Recommended/Not Recommended</i></b><br/>"
        "<i>(By the Head of the Institute)</i><br/><br/><br/>"
        "<i>(Signature of head of institution)</i><br/>"
        "<i>Name & Designation with Seal</i>"
    )
    sign_elem = Paragraph(sign_text, center_st)
            
    t4_data = [
        [photo_elem, "", sign_elem]
    ]
    
    t4 = Table(t4_data, colWidths=[usable_w*0.3, usable_w*0.25, usable_w*0.45])
    t4.setStyle(TableStyle([
        ('VALIGN', (0,0), (-1,-1), 'TOP'),
        ('ALIGN', (2,0), (2,-1), 'CENTER'),
        ('ALIGN', (0,0), (0,-1), 'LEFT')
    ]))
    story.append(t4)
    
    doc.build(story)
    buf.seek(0)
    return buf


# ── Routes ────────────────────────────────────────────────────────────────────

@app.route("/", methods=["GET", "POST"])
def index():
    if request.method == "POST":
        f = request.form
        photo_url = None
        sign_url = None
        token = uuid.uuid4().hex

        track = f.get("Track")
        level = f.get("Level")
        batch_index_str = f.get("Batch_Index")
        
        c_start = ""
        c_end = ""
        wa_link = ""
        gf_link = ""

        if track and level and batch_index_str and batch_index_str.isdigit():
            b_idx = int(batch_index_str)
            config_col = db_client["config"]
            course_dates_doc = config_col.find_one({"_id": "course_dates"}) or {}
            key = f"{track}_{level}"
            batches = course_dates_doc.get(key, [])
            if isinstance(batches, dict):
                batches = [batches] if batches.get("start") else []
            if b_idx < len(batches):
                b = batches[b_idx]
                c_start = b.get("start", "")
                c_end = b.get("end", "")
                wa_link = b.get("wa", "")
                gf_link = b.get("gf", "")

        db = get_db()
        
        # Enforce server-side validation for GOT forms (Educational Qualifications)
        if level in ["Basic", "Advanced"]:
            if not f.get("Edu1_Year") or not f.get("Edu1_Degree") or not f.get("Edu1_University"):
                flash("error", "Educational Qualifications are mandatory for GOT forms.")
                return redirect(url_for("index") + "#nomination")

        doc = {
            "token": token, "submitted_at": datetime.now().isoformat(),
            "track": track, "level": level, "batch_index": batch_index_str,
            "title": f.get("Title"), "name": f.get("Name"), "dob": f.get("DOB"), "gender": f.get("Gender"),
            "contact": f.get("Contact_Number"), "email": f.get("Email"), "aadhar": f.get("Aadhar"),
            "native_state": f.get("Native_State"), "district": f.get("District"),
            "organisation": f.get("Organisation"), "department": f.get("Department"), "designation": f.get("Designation"),
            "status": f.get("Status"), "beneficiary_category": f.get("Beneficiary_Category"),
            "institute_address": f.get("Institute_Address"), "institute_contact": f.get("Institute_Contact"), "institute_email": f.get("Institute_Email"),
            "highest_qual": f.get("Highest_Qualification"),
            "edu1_year": f.get("Edu1_Year"), "edu1_degree": f.get("Edu1_Degree"), "edu1_university": f.get("Edu1_University"),
            "edu2_year": f.get("Edu2_Year"), "edu2_degree": f.get("Edu2_Degree"), "edu2_university": f.get("Edu2_University"),
            "edu3_year": f.get("Edu3_Year"), "edu3_degree": f.get("Edu3_Degree"), "edu3_university": f.get("Edu3_University"),
            "exp1_year": f.get("Exp1_Year"), "exp1_area": f.get("Exp1_Area_of_Expertise"), "exp1_centre": f.get("Exp1_Centre"),
            "exp2_year": f.get("Exp2_Year"), "exp2_area": f.get("Exp2_Area_of_Expertise"), "exp2_centre": f.get("Exp2_Centre"),
            "exp3_year": f.get("Exp3_Year"), "exp3_area": f.get("Exp3_Area_of_Expertise"), "exp3_centre": f.get("Exp3_Centre"),
            "prev_fsp": f.get("Previous_FSP_Program"), "prev_fsp_details": f.get("Previous_FSP_Details_1"),
            "course_start_date": c_start, "course_end_date": c_end, "whatsapp_link": wa_link, "google_form_link": gf_link,
            "resource_centre": f.get("Resource_Centre_Name"),
            "photo_url": photo_url, "sign_url": sign_url
        }

        # Generate PDF and upload to Cloudinary
        form_data = row_to_form_data(doc)
        try:
            pdf_buf = generate_pdf(form_data)
            upload_result = cloudinary.uploader.upload(
                pdf_buf.getvalue(),
                resource_type="raw",
                public_id=f"nominations/nomination_{token}.pdf"
            )
            doc["pdf_url"] = upload_result.get("secure_url", "")
        except Exception as e:
            logging.error(f"Cloudinary upload failed: {e}")
            doc["pdf_url"] = ""

        db.insert_one(doc)
        flash("success", "ok")
        return redirect(url_for("success", token=token))

    config_col = db_client["config"]
    course_dates_doc = config_col.find_one({"_id": "course_dates"}) or {}
    available_courses = {}
    for track_level, batches in course_dates_doc.items():
        if track_level == "_id":
            continue
            
        if isinstance(batches, dict):
            batches = [batches] if batches.get("start") else []
            
        valid_batches = []
        for idx, b in enumerate(batches):
            start = b.get("start")
            end = b.get("end")
            wa = b.get("wa", "")
            if start and end:
                formatted = fmt_course_dates(start, end)
                if formatted:
                    valid_batches.append({"index": idx, "formatted": formatted, "wa": wa})
                    
        if valid_batches:
            available_courses[track_level] = valid_batches

    return render_template("index.html", available_courses=available_courses)


@app.route("/success/<token>")
def success(token):
    db = get_db()
    row = db.find_one({"token": token})
    if not row:
        abort(404)
    return render_template("success.html", token=token, name=row.get("name", "Applicant"), whatsapp_link=row.get("whatsapp_link"), google_form_link=row.get("google_form_link"))


@app.route("/download/pdf/<token>")
def download_pdf(token):
    db = get_db()
    row = db.find_one({"token": token})
    if not row:
        abort(404)
    form_data = row_to_form_data(row)
    pdf_buf = generate_pdf(form_data)
    safe_name = re.sub(r"[^a-zA-Z0-9_\-]", "_",
                       row.get("name") or "Nomination")
    return send_file(
        pdf_buf,
        mimetype="application/pdf",
        as_attachment=True,
        download_name=f"Nomination_{safe_name}.pdf",
    )


@app.route("/download/docx/<token>")
def download_docx(token):
    db = get_db()
    row = db.find_one({"token": token})
    if not row:
        abort(404)
    if not DOCX_TEMPLATE.exists():
        abort(500, "DOCX template missing")
    form_data = row_to_form_data(row)
    docx_buf = generate_docx(form_data)
    safe_name = re.sub(r"[^a-zA-Z0-9_\-]", "_", row["name"] or "Nomination")
    return send_file(
        docx_buf,
        mimetype="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        as_attachment=True,
        download_name=f"Nomination_{safe_name}.docx",
    )


# ── Admin routes ──────────────────────────────────────────────────────────────

@app.route("/admin/dates", methods=["GET", "POST"])
def admin_dates():
    if not session.get("admin"):
        return redirect(url_for("admin"))
    
    config_col = db_client["config"]
    courses = {
        "ARVR_Basic": "GOT – AR & VR (Basic)",
        "ARVR_Advanced": "GOT – AR & VR (Advanced)",
        "ARVR_Bootcamp": "Bootcamp – AR & VR",
        "BDDS_Basic": "GOT – Big Data & Data Science (Basic)",
        "BDDS_Advanced": "GOT – Big Data & Data Science (Advanced)",
        "BDDS_Bootcamp": "Bootcamp – Big Data & Data Science"
    }

    if request.method == "POST":
        updates = {}
        for key in courses.keys():
            starts = request.form.getlist(f"{key}_start[]")
            ends = request.form.getlist(f"{key}_end[]")
            was = request.form.getlist(f"{key}_wa[]")
            gfs = request.form.getlist(f"{key}_gf[]")
            batches = []
            for s, e, w, gf_val in zip(starts, ends, was, gfs):
                if s and e:
                    batches.append({"start": s, "end": e, "wa": w, "gf": gf_val})
            updates[key] = batches
        
        config_col.update_one({"_id": "course_dates"}, {"$set": updates}, upsert=True)
        flash("Course dates updated successfully.", "success")
        return redirect(url_for("admin_dates"))

    doc = config_col.find_one({"_id": "course_dates"}) or {}
    return render_template("admin_dates.html", courses=courses, config_dates=doc)

@app.route("/admin", methods=["GET", "POST"])
def admin():
    if request.method == "POST":
        pwd = request.form.get("password", "")
        if pwd == app.config["ADMIN_PASSWORD"]:
            session["admin"] = True
            return redirect(url_for("admin"))
        flash("Wrong password.", "error")
    if not session.get("admin"):
        return render_template("admin_login.html")

    db = get_db()
    rows = list(db.find().sort("submitted_at", -1))
    for r in rows:
        r["id"] = str(r["_id"])
    return render_template("admin.html", rows=rows)


@app.route("/admin/logout")
def admin_logout():
    session.pop("admin", None)
    return redirect(url_for("admin"))


@app.route("/admin/csv")
def admin_csv():
    if not session.get("admin"):
        abort(403)
    db = get_db()
    rows = list(db.find().sort("submitted_at", -1))
    for r in rows:
        r["id"] = str(r["_id"])
    if not rows:
        flash("No submissions yet.", "warning")
        return redirect(url_for("admin"))

    import csv
    buf = io.StringIO()
    for r in rows:
        r.pop("_id", None)
    cols = list(rows[0].keys())
    writer = csv.DictWriter(buf, fieldnames=cols)
    writer.writeheader()
    for row in rows:
        writer.writerow(row)
    buf.seek(0)
    ts = datetime.now().strftime("%Y-%m-%d_%H-%M")
    return send_file(
        io.BytesIO(buf.getvalue().encode("utf-8-sig")),
        mimetype="text/csv",
        as_attachment=True,
        download_name=f"NIELIT_Nominations_{ts}.csv",
    )


@app.route("/admin/pdf/<string:row_id>")
def admin_pdf_single(row_id):
    if not session.get("admin"):
        abort(403)
    db = get_db()
    row = db.find_one({"_id": ObjectId(row_id)})
    if not row:
        abort(404)
    form_data = row_to_form_data(row)
    pdf_buf = generate_pdf(form_data)
    safe_name = re.sub(r"[^a-zA-Z0-9_\-]", "_",
                       row.get("name") or "Nomination")
    return send_file(pdf_buf, mimetype="application/pdf", as_attachment=True,
                     download_name=f"Nomination_{safe_name}.pdf")


@app.route("/admin/pdf/all")
def admin_pdf_all():
    if not session.get("admin"):
        abort(403)
    db = get_db()
    rows = list(db.find().sort("submitted_at", 1))
    if not rows:
        flash("No submissions yet.", "warning")
        return redirect(url_for("admin"))

    writer = PdfWriter()
    for row in rows:
        form_data = row_to_form_data(row)
        pdf_buf = generate_pdf(form_data)
        reader = PdfReader(pdf_buf)
        for page in reader.pages:
            writer.add_page(page)

    out_buf = BytesIO()
    writer.write(out_buf)
    out_buf.seek(0)
    ts = datetime.now().strftime("%Y-%m-%d_%H-%M")
    return send_file(out_buf, mimetype="application/pdf", as_attachment=True,
                     download_name=f"All_Nominations_{ts}.pdf")



@app.route("/admin/docx/all")
def admin_docx_all():
    if not session.get("admin"):
        abort(403)
    db = get_db()
    rows = list(db.find().sort("submitted_at", 1))
    if not rows:
        flash("No submissions yet.", "warning")
        return redirect(url_for("admin"))

    master = None
    for row in rows:
        form_data = row_to_form_data(row)
        docx_buf = generate_docx(form_data)
        doc = Document(docx_buf)
        if master is None:
            master = Composer(doc)
        else:
            master.append(doc)

    out_buf = BytesIO()
    master.save(out_buf)
    out_buf.seek(0)
    ts = datetime.now().strftime("%Y-%m-%d_%H-%M")
    return send_file(out_buf, mimetype="application/vnd.openxmlformats-officedocument.wordprocessingml.document", as_attachment=True,
                     download_name=f"All_Nominations_{ts}.docx")


@app.route("/admin/delete_all", methods=["POST"])
def admin_delete_all():
    if not session.get("admin"):
        abort(403)
        
    pwd = request.form.get("delete_password", "")
    expected_pwd = os.environ.get("DELETE_ALL_PASSWORD", "rccrcc")
    if pwd != expected_pwd:
        flash("Incorrect password to delete all.", "error")
        return redirect(url_for("admin"))

    db = get_db()
    db.delete_many({})
    flash("All nominations deleted successfully.", "success")
    return redirect(url_for("admin"))

@app.route("/admin/zip/pdfs")
def admin_zip_pdfs():
    if not session.get("admin"):
        abort(403)
    db = get_db()
    rows = list(db.find())
    if not rows:
        flash("No submissions yet.", "warning")
        return redirect(url_for("admin"))
    
    zip_buf = BytesIO()
    with zipfile.ZipFile(zip_buf, "w", zipfile.ZIP_DEFLATED) as zf:
        for row in rows:
            form_data = row_to_form_data(row)
            pdf_buf = generate_pdf(form_data)
            safe_name = re.sub(r"[^a-zA-Z0-9_\-]", "_", row.get("name") or "Nomination")
            zf.writestr(f"Nomination_{safe_name}_{row.get('token')[:6]}.pdf", pdf_buf.getvalue())
            
    zip_buf.seek(0)
    ts = datetime.now().strftime("%Y-%m-%d_%H-%M")
    return send_file(zip_buf, mimetype="application/zip", as_attachment=True, download_name=f"All_PDFs_{ts}.zip")

@app.route("/admin/zip/docxs")
def admin_zip_docxs():
    if not session.get("admin"):
        abort(403)
    db = get_db()
    rows = list(db.find())
    if not rows:
        flash("No submissions yet.", "warning")
        return redirect(url_for("admin"))
    
    zip_buf = BytesIO()
    with zipfile.ZipFile(zip_buf, "w", zipfile.ZIP_DEFLATED) as zf:
        for row in rows:
            form_data = row_to_form_data(row)
            docx_buf = generate_docx(form_data)
            safe_name = re.sub(r"[^a-zA-Z0-9_\-]", "_", row.get("name") or "Nomination")
            zf.writestr(f"Nomination_{safe_name}_{row.get('token')[:6]}.docx", docx_buf.getvalue())
            
    zip_buf.seek(0)
    ts = datetime.now().strftime("%Y-%m-%d_%H-%M")
    return send_file(zip_buf, mimetype="application/zip", as_attachment=True, download_name=f"All_DOCXs_{ts}.zip")


@app.route("/admin/delete/<string:row_id>", methods=["POST"])
def admin_delete(row_id):
    if not session.get("admin"):
        abort(403)
    db = get_db()
    try:
        obj_id = ObjectId(row_id)
    except Exception:
        abort(400)
    row = db.find_one({"_id": obj_id})
    if row:
        db.delete_one({"_id": obj_id})
    flash("Entry deleted.", "success")
    return redirect(url_for("admin"))


# ── Init & run ────────────────────────────────────────────────────────────────

@app.errorhandler(404)
def page_not_found(e):
    return render_template('404.html'), 404


if __name__ == "__main__":
    port = int(os.environ.get("PORT", 5000))
    app.run(host="0.0.0.0", port=port, debug=True)
