
import os
import re
import logging
from io import BytesIO
from datetime import datetime
from pathlib import Path
from docx import Document
from reportlab.platypus import (Image as RLImage, Paragraph, SimpleDocTemplate, Spacer, Table, TableStyle)
from reportlab.lib.units import cm, mm, inch
from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
from reportlab.lib.pagesizes import A4
from reportlab.lib import colors

from services.helpers import fmt_date, fmt_course_dates, _course_name, _technology
from services.media import fetch_image_as_jpeg
from models.database import get_config_col

DOCX_TEMPLATE = Path('docxtemplates/GOT_Nomination_Form.docx')

def row_to_form_data(row):
    """Map DB row to the placeholder dict used by DOCX/PDF generation."""
    d = dict(row)
    name_full = f"{d.get('title','')} {d.get('name','')}".strip()
    
    track = (d.get("track") or "").upper()
    level = (d.get("level") or "").capitalize()
    course_key = f"{track}_{level}"
    
    c_start = d.get("course_start_date")
    c_end = d.get("course_end_date")
    
    if not c_start or not c_end:
        # Fallback for old nominations that didn't save dates directly
        config_col = get_config_col()
        course_dates_doc = config_col.find_one({"_id": "course_dates"}) or {}
        batches = course_dates_doc.get(course_key, [])
        if isinstance(batches, dict):
            batches = [batches] if batches.get("start") else []
        if batches:
            c_start = c_start or batches[0].get("start")
            c_end = c_end or batches[0].get("end")

    formatted_dates = fmt_course_dates(c_start, c_end) if c_start and c_end else ""
    
    return {
        "Title": d.get("title", ""),
        "Name": d.get("name", ""),
        "Full_Name": name_full,
        "Course_Name": _course_name(d.get("track", ""), d.get("level", "")),
        "Course_Start_Date": formatted_dates,
        "Date_of_Training": formatted_dates,
        "Native_State": d.get("native_state", ""),
        "District": d.get("district", ""),
        "Status": d.get("status", ""),
        "Beneficiary_Category": d.get("beneficiary_category", ""),
        "Level": d.get("level", ""),
        "Applicant_Name": name_full,
        "Gov_ID_Number": d.get("aadhar", ""),
        "Organization_Academic_Institute": d.get("organisation", ""),
        "Role": d.get("designation", ""),
        "Highest_Qualification": f"{d.get('highest_qual', '')}",
        "Previous_FSP_Program": d.get("prev_fsp", ""),
        "Previous_FSP_Details_1": d.get("prev_fsp_details", ""),
        "Previous_FSP_Details_2": "",
        "Technology": _technology(d.get("track", "")),
        "Resource_Centre_Name": d.get("resource_centre", "NIELIT Chandigarh"),
        "Date_of_Training": formatted_dates,
        "Designation": d.get("designation", ""),
        "Organisation": d.get("organisation", ""),
        "DOB": fmt_date(d.get("dob", "")),
        "Gender": d.get("gender", ""),
        "Aadhar": d.get("aadhar", ""),
        "Contact_Number": d.get("contact", ""),
        "Email": d.get("email", ""),
        "Organisation_Department": _merge(d.get("organisation", ""), d.get("department", "")),
        "Contact_Number_Email": _merge(d.get("contact", ""), d.get("email", "")),
        "Institute_Details": _merge(_merge(d.get("institute_address", ""), d.get("institute_contact", "")), d.get("institute_email", "")),
        "photo_url": d.get("photo_url", ""),
        "sign_url": d.get("sign_url", ""),
        "Edu1_Year": d.get("edu1_year", ""), "Edu1_Degree": d.get("edu1_degree", ""), "Edu1_University": d.get("edu1_university", ""),
        "Edu2_Year": d.get("edu2_year", ""), "Edu2_Degree": d.get("edu2_degree", ""), "Edu2_University": d.get("edu2_university", ""),
        "Edu3_Year": d.get("edu3_year", ""), "Edu3_Degree": d.get("edu3_degree", ""), "Edu3_University": d.get("edu3_university", ""),
        "Exp1_Year": d.get("exp1_year", ""), "Exp1_Area_of_Expertise": d.get("exp1_area", ""), "Exp1_Centre": d.get("exp1_centre", ""),
        "Exp2_Year": d.get("exp2_year", ""), "Exp2_Area_of_Expertise": d.get("exp2_area", ""), "Exp2_Centre": d.get("exp2_centre", ""),
        "Exp3_Year": d.get("exp3_year", ""), "Exp3_Area_of_Expertise": d.get("exp3_area", ""), "Exp3_Centre": d.get("exp3_centre", ""),
        "Today_Date": datetime.today().strftime("%d-%m-%Y"),
        "photo_url": d.get("photo_url"),
        "sign_url": d.get("sign_url"),
    }


def _course_name(track, level):
    track = (track or "").upper()
    level = (level or "").capitalize()
    names = {
        ("ARVR", "Basic"):    "GOT – AR & VR (Basic)",
        ("ARVR", "Advanced"): "GOT – AR & VR (Advanced)",
        ("ARVR", "Bootcamp"): "Bootcamp – AR & VR",
        ("BDDS", "Basic"):    "GOT – Big Data & Data Science (Basic)",
        ("BDDS", "Advanced"): "GOT – Big Data & Data Science (Advanced)",
        ("BDDS", "Bootcamp"): "Bootcamp – Big Data & Data Science",
    }
    return names.get((track, level), f"{track} {level}")


def _technology(track):
    t = (track or "").upper()
    return {"ARVR": "Augmented & Virtual Reality", "BDDS": "Big Data & Data Science"}.get(t, track)


def _merge(a, b, sep=" / "):
    parts = [x for x in [a, b] if x]
    return sep.join(parts) if parts else ""


# ── DOCX generation ───────────────────────────────────────────────────────────
def _replace_in_para(para, replacements):
    full = para.text
    if "{" not in full and "<<" not in full:
        return
    if para.runs:
        for run in para.runs:
            for k, v in replacements.items():
                run.text = run.text.replace(
                    f"{{{k}}}", v).replace(f"<<{k}>>", v)
        joined = "".join(r.text for r in para.runs)
        if "{" in joined or "<<" in joined:
            updated = joined
            for k, v in replacements.items():
                updated = updated.replace(f"{{{k}}}", v).replace(f"<<{k}>>", v)
            if updated != joined:
                para.runs[0].text = updated
                for r in para.runs[1:]:
                    r.text = ""


def generate_docx(form_data: dict) -> BytesIO:
    """Fill the DOCX template and return as BytesIO. Inserts photo into Photo cell."""
    template_path = Path("docxtemplates/Bootcamp_Nomination_Form.docx") if form_data.get("Level") == "Bootcamp" else DOCX_TEMPLATE
    doc = Document(str(template_path))
    
    # Squeeze margins slightly to guarantee everything fits after inserting logo
    from docx.shared import Inches
    for section in doc.sections:
        if section.bottom_margin > Inches(0.25):
            section.bottom_margin = Inches(0.25)
        if section.top_margin > Inches(0.25):
            section.top_margin = Inches(0.25)

    def add_constrained_picture(run, img_path, max_width_inch, max_height_inch):
        from PIL import Image
        with Image.open(img_path) as im:
            w, h = im.size
        aspect = w / float(h) if h > 0 else 1
        target_aspect = max_width_inch / float(max_height_inch)
        if aspect > target_aspect:
            run.add_picture(img_path, width=Inches(max_width_inch))
        else:
            run.add_picture(img_path, height=Inches(max_height_inch))

    replacements = {k: (str(v) if v else "") for k, v in form_data.items()
                    if k not in ("photo_url", "sign_url")}

    level_val = (form_data.get("Level") or "").capitalize()
    if level_val in ("Advanced", "Basic"):
        prog_title = f"Government Officials Training ({level_val}) Programme"
    else:
        prog_title = "Government Officials Training Programme"

    for para in doc.paragraphs:
        if "Government Officials Training Programme" in para.text:
            for run in para.runs:
                run.text = run.text.replace("Government Officials Training Programme", prog_title)

        cell_text = para.text.strip()
        if "{Photo}" in cell_text or "<<Photo>>" in cell_text or cell_text == "Photo":
            para.text = ""
            if form_data.get("photo_url"):
                try:
                    img_path = fetch_image_as_jpeg(form_data["photo_url"], target_size=(400, 400))
                    add_constrained_picture(para.add_run(), img_path, 1.1, 1.3)
                    os.remove(img_path)
                except Exception as e:
                    logging.error(f"Failed to add photo to docx para: {e}")
                    para.text = "Photo"
            else:
                para.text = "Photo"
            continue
        elif "{Signature}" in cell_text or "<<Signature>>" in cell_text or cell_text == "Signature":
            para.text = ""
            if form_data.get("sign_url"):
                try:
                    img_path = fetch_image_as_jpeg(form_data["sign_url"], target_size=(600, 200))
                    add_constrained_picture(para.add_run(), img_path, 1.5, 0.5)
                    os.remove(img_path)
                except Exception as e:
                    logging.error(f"Failed to add sign to docx para: {e}")
                    para.text = "[Please Sign Here]"
            else:
                para.text = "[Please Sign Here]"
            continue
        
    photo_added = False
    sign_added = False
    import re

    # Wiping cells that contain 'Photo' ensures we delete any margin-heavy Text Boxes and insert the image natively
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                if not photo_added and re.search(r'>\s*\{?Photo\}?\s*<', cell._element.xml, re.IGNORECASE):
                    cell.text = "" 
                    if form_data.get("photo_url"):
                        try:
                            img_path = fetch_image_as_jpeg(form_data["photo_url"], target_size=(250, 250))
                            if not cell.paragraphs:
                                cell.add_paragraph()
                            p = cell.paragraphs[0]
                            p.alignment = 1
                            add_constrained_picture(p.add_run(), img_path, 1.0, 1.2)
                            os.remove(img_path)
                        except Exception as e:
                            logging.error(f"Photo error: {e}")
                    photo_added = True

    def process_paragraph(para):
        nonlocal photo_added, sign_added
        
        # Replace image placeholders
        ptxt = para.text.strip()
        if not photo_added and ("{Photo}" in ptxt or "<<Photo>>" in ptxt or ptxt == "Photo"):
            para.text = ""
            if form_data.get("photo_url"):
                try:
                    img_path = fetch_image_as_jpeg(form_data["photo_url"], target_size=(250, 250))
                    add_constrained_picture(para.add_run(), img_path, 1.0, 1.2)
                    os.remove(img_path)
                except Exception as e:
                    logging.error(f"Photo error: {e}")
            photo_added = True
            
        elif not sign_added and ("{Signature}" in ptxt or "<<Signature>>" in ptxt or ptxt == "Signature"):
            para.text = ""
            if form_data.get("sign_url"):
                try:
                    img_path = fetch_image_as_jpeg(form_data["sign_url"], target_size=(300, 100))
                    add_constrained_picture(para.add_run(), img_path, 1.5, 0.5)
                    os.remove(img_path)
                except Exception as e:
                    logging.error(f"Sign error: {e}")
            sign_added = True

        elif not sign_added and ("Signature of the Official" in ptxt or "Applicant Signature" in ptxt):
            if form_data.get("sign_url"):
                try:
                    from docx.shared import Pt
                    img_path = fetch_image_as_jpeg(form_data["sign_url"], target_size=(300, 100))
                    new_para = para.insert_paragraph_before("")
                    new_para.alignment = 2 # Right align
                    new_para.paragraph_format.space_after = Pt(0)
                    new_para.paragraph_format.space_before = Pt(0)
                    add_constrained_picture(new_para.add_run(), img_path, 1.5, 0.5)
                    os.remove(img_path)
                except Exception as e:
                    logging.error(f"Sign error: {e}")
            sign_added = True

        # Handle formatting hacks
        if "Any of the government issued ID" in para.text:
            para.style = 'Normal'
            para.text = ""
            r0 = para.add_run("• ")
            r1 = para.add_run("1")
            r1.font.superscript = True
            r2 = para.add_run(" Any of the government issued ID: Aadhaar card")
            r2.font.italic = True
        elif "Participants are not allowed to enroll" in para.text:
            para.style = 'Normal'
            para.text = ""
            r0 = para.add_run("• ")
            r1 = para.add_run("2")
            r1.font.superscript = True
            r2 = para.add_run(" Participants are not allowed to enroll in the same program for more than once")
            r2.font.italic = True
        else:
            _replace_in_para(para, replacements)

    from docx.text.paragraph import Paragraph
    all_paras = [Paragraph(p, doc) for p in doc._element.xpath('.//w:p')]
    for para in all_paras:
        process_paragraph(para)

    for section in doc.sections:
        for para in section.header.paragraphs:
            _replace_in_para(para, replacements)
        for para in section.footer.paragraphs:
            _replace_in_para(para, replacements)

    # Fallback: Add Photo and Signature at the end if they were NOT added above but exist
    needs_photo_fallback = not photo_added and form_data.get("photo_url")
    needs_sign_fallback = not sign_added and form_data.get("sign_url")
    
    if needs_photo_fallback or needs_sign_fallback:
        doc.add_paragraph()
        table = doc.add_table(rows=1, cols=3)
        if needs_photo_fallback:
            try:
                img_path = fetch_image_as_jpeg(form_data["photo_url"], target_size=(400, 400))
                add_constrained_picture(table.cell(0, 0).paragraphs[0].add_run(), img_path, 1.1, 1.3)
                os.remove(img_path)
            except: pass
        if needs_sign_fallback:
            try:
                img_path = fetch_image_as_jpeg(form_data["sign_url"], target_size=(600, 200))
                p = table.cell(0, 2).paragraphs[0]
                p.alignment = 1
                add_constrained_picture(p.add_run(), img_path, 1.5, 0.5)
                os.remove(img_path)
            except: pass

    buf = BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf


# ── PDF generation via ReportLab ──────────────────────────────────────────────
def generate_pdf(form_data: dict) -> BytesIO:
    """Generate a simple A4 PDF nomination form matching the DOCX template."""
    buf = BytesIO()
    
    if form_data.get("Level") == "Bootcamp":
        doc = SimpleDocTemplate(
            buf, pagesize=A4,
            topMargin=0.8*cm, bottomMargin=0.8*cm,
            leftMargin=1.2*cm, rightMargin=1.2*cm,
            title="Nomination Form",
            author="FutureSkills PRIME"
        )
        elements = []
        
        # Styles
        styles = getSampleStyleSheet()
        title_st = ParagraphStyle('title', fontName='Helvetica-Bold', fontSize=11, alignment=1, spaceAfter=2)
        norm_st = ParagraphStyle('norm', fontName='Helvetica', fontSize=8)
        bold_st = ParagraphStyle('bold', fontName='Helvetica-Bold', fontSize=8)
        
        usable_w = A4[0] - 2.4*cm
        
        # Header with Logos
        try:
            with open("static/assets/img/NIELIT-Logo-hd.png", "rb") as f:
                n_data = BytesIO(f.read())
            nielit_img = RLImage(n_data, width=1.3*inch, height=0.7*inch)
        except Exception:
            nielit_img = ""
            
        try:
            fs_img = RLImage("static/assets/img/FUTURESKILLS_LOGO.png", width=1.5*inch, height=0.7*inch)
        except Exception:
            fs_img = ""
            
        header_text = """
        <font color="#000080" size="14"><b>National Institute of Electronics and<br/>Information Technology Chandigarh</b></font><br/>
        <font color="#000080" size="7">An autonomous scientific society under administrative control of<br/>
        Ministry of Electronics and Information Technology (MeitY), Government of India</font>
        """
        header_p = Paragraph(header_text, ParagraphStyle('h', alignment=1, leading=10))
        
        header_t = Table([[nielit_img, header_p, fs_img]], colWidths=[1.5*inch, usable_w-3.0*inch, 1.5*inch])
        header_t.setStyle(TableStyle([
            ('VALIGN', (0,0), (-1,-1), 'MIDDLE'),
            ('ALIGN', (0,0), (-1,-1), 'CENTER'),
        ]))
        elements.append(header_t)
        
        fsp_text = "<b>FutureSkills PRIME (Programme for Re-skilling/Up-skilling of ITManpower for Employability)</b>"
        elements.append(Paragraph(fsp_text, ParagraphStyle('fsp', fontName='Helvetica-Bold', fontSize=9, alignment=1)))
        elements.append(Spacer(1, 10))
        
        elements.append(Paragraph("<b>Application Form for Bootcamp</b>", ParagraphStyle('afb', fontName='Helvetica-Bold', fontSize=12, alignment=1)))
        elements.append(Spacer(1, 10))
        
        elements.append(Paragraph("<b>Date:</b>", ParagraphStyle('rdate', fontName='Helvetica-Bold', fontSize=9, alignment=2, rightIndent=100)))
        elements.append(Spacer(1, 5))
        
        data = [
            [Paragraph("<b>Resource Centre Name</b>", norm_st), "", ":", Paragraph(form_data.get("Resource_Centre_Name", ""), norm_st), ""],
            [Paragraph("<b>Technology</b>", norm_st), "", ":", Paragraph(form_data.get("Technology", ""), norm_st), Paragraph("<b>Role:</b> Co-Lead", norm_st)],
            ["1", Paragraph("Course Name", norm_st), ":", Paragraph(form_data.get("Course_Name", ""), norm_st), ""],
            ["2", Paragraph("Course Start Date", norm_st), ":", Paragraph(form_data.get("Course_Start_Date", ""), norm_st), ""],
            ["3", Paragraph("Applicant Name (as per Gov ID)", norm_st), ":", Paragraph(form_data.get("Applicant_Name", ""), norm_st), ""],
            ["4", Paragraph("Date of Birth", norm_st), ":", Paragraph(form_data.get("DOB", ""), norm_st), ""],
            ["5", Paragraph("Gender", norm_st), ":", Paragraph(form_data.get("Gender", ""), norm_st), ""],
            ["6", Paragraph("Mobile Number", norm_st), ":", Paragraph(form_data.get("Contact_Number", ""), norm_st), ""],
            ["7", Paragraph("Email ID (official ID preferred)", norm_st), ":", Paragraph(form_data.get("Email", ""), norm_st), ""],
            ["8", Paragraph("Native State", norm_st), ":", Paragraph(form_data.get("Native_State", ""), norm_st), ""],
            ["9", Paragraph("District", norm_st), ":", Paragraph(form_data.get("District", ""), norm_st), ""],
            ["10", Paragraph("Government-issued ID Number<sup>1</sup><br/><i>(Aadhar copy enclosed)</i>", norm_st), ":", Paragraph(form_data.get("Gov_ID_Number", ""), norm_st), ""],
            ["11", Paragraph("Organization/Academic Institute (if applicable)", norm_st), ":", Paragraph(form_data.get("Organization_Academic_Institute", ""), norm_st), ""],
            ["12", Paragraph("Highest Qualification (with Degree & Branch)", norm_st), ":", Paragraph(form_data.get("Highest_Qualification", ""), norm_st), ""],
            ["13", Paragraph("Status (Pursuing/ Passed out)", norm_st), ":", Paragraph(form_data.get("Status", ""), norm_st), ""],
            ["14", Paragraph("Beneficiary Category (Tick as applicable)", norm_st), ":", Paragraph(form_data.get("Beneficiary_Category", ""), norm_st), ""],
            ["15", Paragraph("Involved in previous FSP Program<sup>2</sup>", norm_st), ":", Paragraph(form_data.get("Previous_FSP_Program", ""), norm_st), ""],
            ["16", Paragraph("If previous answer is, yes provide details<br/>(Program Name / Conducting Institute / Date)", norm_st), ":", Paragraph(f"1. {form_data.get('Previous_FSP_Details_1', '')}<br/>2. {form_data.get('Previous_FSP_Details_2', '')}", norm_st), ""],
        ]
        
        t_style_cmds = [
            ('GRID', (0,2), (-1,-1), 0.5, colors.black),
            ('LINEABOVE', (0,0), (-1,0), 0.5, colors.black),
            ('LINEBELOW', (0,0), (-1,0), 0.5, colors.black),
            ('LINEBELOW', (0,1), (-1,1), 0.5, colors.black),
            ('VALIGN', (0,0), (-1,-1), 'MIDDLE'),
            ('PADDING', (0,0), (-1,-1), 2),
            ('SPAN', (0, 0), (1, 0)), 
            ('SPAN', (3, 0), (4, 0)), 
            ('SPAN', (0, 1), (1, 1)), 
        ]
        
        for r in range(2, len(data)):
            t_style_cmds.append(('SPAN', (3, r), (4, r)))
            
        t_style = TableStyle(t_style_cmds)
        
        t = Table(data, colWidths=[usable_w*0.05, usable_w*0.35, usable_w*0.05, usable_w*0.35, usable_w*0.20])
        t.setStyle(t_style)
        elements.append(t)
        elements.append(Spacer(1, 5))
        
        elements.append(Paragraph("I hereby declare that all the information provided above is true.", norm_st))
        
        elements.append(Spacer(1, 10))
        
        photo_elem = ""
        if form_data.get('photo_url'):
            try:
                img_path = fetch_image_as_jpeg(form_data['photo_url'], target_size=(400, 400))
                photo_elem = RLImage(img_path, width=1.1*inch, height=1.1*inch)
            except:
                pass
                
        sign_elem = [Paragraph("<b>Applicant Signature</b>", ParagraphStyle('r', fontName='Helvetica', fontSize=8, alignment=2))]
        if form_data.get('sign_url'):
            try:
                img_path = fetch_image_as_jpeg(form_data['sign_url'], target_size=(600, 200))
                sign_elem.insert(0, RLImage(img_path, width=1.5*inch, height=0.5*inch, hAlign='RIGHT'))
            except:
                pass
                
        bot_table = Table([[photo_elem, sign_elem]], colWidths=[usable_w*0.5, usable_w*0.5])
        bot_table.setStyle(TableStyle([
            ('ALIGN', (0,0), (0,0), 'LEFT'),
            ('ALIGN', (1,0), (1,0), 'RIGHT'),
            ('VALIGN', (0,0), (-1,-1), 'BOTTOM'),
        ]))
        elements.append(bot_table)
        
        elements.append(Spacer(1, 5))
        long_line = Table([[""]], colWidths=[usable_w], rowHeights=[1])
        long_line.setStyle(TableStyle([('LINEABOVE', (0,0), (-1,-1), 0.5, colors.black)]))
        elements.append(long_line)
        
        elements.append(Spacer(1, 5))
        elements.append(Paragraph("<b>(For office purpose)</b><br/><br/>The above submitted information has been verified and recommended.", norm_st))
        
        elements.append(Spacer(1, 10))
        elements.append(Paragraph("(Signature)<br/><b>Course Co-ordinator</b>", ParagraphStyle('r2', fontName='Helvetica', fontSize=8, alignment=2)))
        
        elements.append(Spacer(1, 5))
        short_line = Table([[""]], colWidths=[150], rowHeights=[1], hAlign='LEFT')
        short_line.setStyle(TableStyle([('LINEABOVE', (0,0), (-1,-1), 0.5, colors.black)]))
        elements.append(short_line)
        
        elements.append(Spacer(1, 2))
        bullet_text = "&bull; &nbsp; <sup>1</sup> <i>Any of the government issued ID: Aadhaar card</i><br/>" \
                      "&bull; &nbsp; <sup>2</sup> <i>Participants are not allowed to enroll in the same program for more than once</i>"
        elements.append(Paragraph(bullet_text, ParagraphStyle('small', fontName='Helvetica', fontSize=7)))
        
        def draw_page_border(canvas, document):
            canvas.saveState()
            canvas.setStrokeColor(colors.black)
            canvas.setLineWidth(3)
            canvas.rect(0.6*cm, 0.6*cm, A4[0]-1.2*cm, A4[1]-1.2*cm)
            canvas.setLineWidth(0.5)
            canvas.rect(0.65*cm, 0.65*cm, A4[0]-1.3*cm, A4[1]-1.3*cm)
            canvas.restoreState()
            
        doc.build(elements, onFirstPage=draw_page_border, onLaterPages=draw_page_border)
        buf.seek(0)
        return buf
        
    # --- END BOOTCAMP BRANCH ---


    doc = SimpleDocTemplate(
        buf, pagesize=A4,
        topMargin=0.5*cm, bottomMargin=0.5*cm,
        leftMargin=1.2*cm, rightMargin=1.2*cm,
        title="Nomination Form",
        author="FutureSkills PRIME"
    )
    W, H = A4
    usable_w = W - 3.6*cm

    styles = getSampleStyleSheet()
    
    title_st = ParagraphStyle('title', fontName='Helvetica-Bold', fontSize=11, alignment=1, spaceAfter=2)
    sub_st = ParagraphStyle('sub', fontName='Helvetica-Bold', fontSize=9, alignment=1, spaceAfter=4)
    norm_st = ParagraphStyle('norm', fontName='Helvetica', fontSize=8)
    bold_st = ParagraphStyle('bold', fontName='Helvetica-Bold', fontSize=8)
    right_st = ParagraphStyle('right', fontName='Helvetica', fontSize=8, alignment=2)
    center_st = ParagraphStyle('center_st', fontName='Helvetica', fontSize=9, alignment=1)
    
    story = []
    
    story.append(Paragraph("Nomination Form", title_st))
    
    level_val = (form_data.get("Level") or "").capitalize()
    if level_val in ("Advanced", "Basic"):
        prog_title = f"Government Officials Training ({level_val}) Programme"
    else:
        prog_title = "Government Officials Training Programme"
        
    story.append(Paragraph(prog_title, sub_st))
    
    # Table 0: Training Programme Details
    story.append(Paragraph("Training Programme Details", bold_st))
    t0_data = [
        [Paragraph("Name of Training Programme", norm_st), Paragraph(form_data.get("Course_Name", ""), norm_st)],
        [Paragraph("Name of Technology", norm_st), Paragraph(form_data.get("Technology", ""), norm_st)],
        [Paragraph("Resource Centre Name", norm_st), Paragraph(form_data.get("Resource_Centre_Name", ""), norm_st)],
        [Paragraph("Date of Training", norm_st), Paragraph(form_data.get("Date_of_Training", ""), norm_st)]
    ]
    t0 = Table(t0_data, colWidths=[usable_w*0.4, usable_w*0.6])
    t0.setStyle(TableStyle([
        ('BOX', (0,0), (-1,-1), 0.5, colors.black),
        ('INNERGRID', (0,0), (-1,-1), 0.5, colors.black),
        ('VALIGN', (0,0), (-1,-1), 'MIDDLE'),
        ('PADDING', (0,0), (-1,-1), 1)
    ]))
    story.append(t0)
    story.append(Spacer(1, 2))
    
    # Table 1: Personal Information
    story.append(Paragraph("Personal Information", bold_st))
    t1_data = [
        [Paragraph("NAME<br/>Prof./Dr./Mr./Ms.", norm_st), Paragraph(f"{form_data.get('Title','')} {form_data.get('Name','')}", norm_st), "", ""],
        [Paragraph("DESIGNATION:", norm_st), Paragraph(form_data.get("Designation", ""), norm_st), Paragraph("ORGANISATION:", norm_st), Paragraph(form_data.get("Organisation", ""), norm_st)],
        [Paragraph("DATE OF BIRTH :", norm_st), Paragraph(form_data.get("DOB", ""), norm_st), Paragraph("GENDER(M/F)", norm_st), Paragraph(form_data.get("Gender", ""), norm_st)],
        [Paragraph("AADHAR No.", norm_st), Paragraph(form_data.get("Aadhar", ""), norm_st), "", ""],
        [Paragraph("CONTACT NUMBER &<br/>E- MAIL", norm_st), Paragraph(f"{form_data.get('Contact_Number','')} / {form_data.get('Email','')}", norm_st), "", ""],
        [Paragraph("NAME OF THE<br/>ORGANISATION/<br/>DEPARTMENT", norm_st), Paragraph(form_data.get("Organisation_Department", ""), norm_st), "", ""],
        [Paragraph("COMPLETE ADDRESS /<br/>CONTACT NUMBERS /<br/>E- MAIL OF THE INSTITUTE", norm_st), Paragraph(form_data.get('Institute_Details', ''), norm_st), "", ""]
    ]
    t1 = Table(t1_data, colWidths=[usable_w*0.25, usable_w*0.25, usable_w*0.25, usable_w*0.25])
    t1.setStyle(TableStyle([
        ('BOX', (0,0), (-1,-1), 0.5, colors.black),
        ('INNERGRID', (0,0), (-1,-1), 0.5, colors.black),
        ('SPAN', (1,0), (3,0)),
        ('SPAN', (1,3), (3,3)),
        ('SPAN', (1,4), (3,4)),
        ('SPAN', (1,5), (3,5)),
        ('SPAN', (1,6), (3,6)),
        ('VALIGN', (0,0), (-1,-1), 'MIDDLE'),
        ('PADDING', (0,0), (-1,-1), 1)
    ]))
    story.append(t1)
    story.append(Spacer(1, 2))
    
    # Table 2: Educational Qualifications
    story.append(Paragraph("Educational / Professional Qualifications", bold_st))
    t2_data = [
        [Paragraph("EDUCATIONAL / PROFESSIONAL QUALIFICATIONS (GRADUATION ONWARDS)", bold_st), "", "", ""],
        [Paragraph("SL. No.", bold_st), Paragraph("YEAR", bold_st), Paragraph("DEGREE", bold_st), Paragraph("UNIVERSITY/INSTITUTE", bold_st)],
        ["1", Paragraph(form_data.get("Edu1_Year",""), norm_st), Paragraph(form_data.get("Edu1_Degree",""), norm_st), Paragraph(form_data.get("Edu1_University",""), norm_st)],
        ["2", Paragraph(form_data.get("Edu2_Year",""), norm_st), Paragraph(form_data.get("Edu2_Degree",""), norm_st), Paragraph(form_data.get("Edu2_University",""), norm_st)],
        ["3", Paragraph(form_data.get("Edu3_Year",""), norm_st), Paragraph(form_data.get("Edu3_Degree",""), norm_st), Paragraph(form_data.get("Edu3_University",""), norm_st)]
    ]
    t2 = Table(t2_data, colWidths=[usable_w*0.1, usable_w*0.2, usable_w*0.3, usable_w*0.4])
    t2.setStyle(TableStyle([
        ('BOX', (0,0), (-1,-1), 0.5, colors.black),
        ('INNERGRID', (0,0), (-1,-1), 0.5, colors.black),
        ('SPAN', (0,0), (3,0)),
        ('VALIGN', (0,0), (-1,-1), 'MIDDLE'),
        ('ALIGN', (0,0), (3,1), 'CENTER'),
        ('PADDING', (0,0), (-1,-1), 1)
    ]))
    story.append(t2)
    story.append(Spacer(1, 2))
    
    # Table 3: Experience
    t3_data = [
        [Paragraph("RESEARCH / TECHNICAL EXPERIENCE", bold_st), "", "", ""],
        [Paragraph("SL.NO.", bold_st), Paragraph("YEAR", bold_st), Paragraph("AREA OF EXPERTISE", bold_st), Paragraph("CENTRE", bold_st)],
        ["1", Paragraph(form_data.get("Exp1_Year",""), norm_st), Paragraph(form_data.get("Exp1_Area_of_Expertise",""), norm_st), Paragraph(form_data.get("Exp1_Centre",""), norm_st)],
        ["2", Paragraph(form_data.get("Exp2_Year",""), norm_st), Paragraph(form_data.get("Exp2_Area_of_Expertise",""), norm_st), Paragraph(form_data.get("Exp2_Centre",""), norm_st)],
        ["3", Paragraph(form_data.get("Exp3_Year",""), norm_st), Paragraph(form_data.get("Exp3_Area_of_Expertise",""), norm_st), Paragraph(form_data.get("Exp3_Centre",""), norm_st)]
    ]
    t3 = Table(t3_data, colWidths=[usable_w*0.1, usable_w*0.2, usable_w*0.4, usable_w*0.3])
    t3.setStyle(TableStyle([
        ('BOX', (0,0), (-1,-1), 0.5, colors.black),
        ('INNERGRID', (0,0), (-1,-1), 0.5, colors.black),
        ('SPAN', (0,0), (3,0)),
        ('VALIGN', (0,0), (-1,-1), 'MIDDLE'),
        ('ALIGN', (0,0), (3,1), 'CENTER'),
        ('PADDING', (0,0), (-1,-1), 1)
    ]))
    story.append(t3)
    story.append(Spacer(1, 2))
    
    # Table 4: Signature / Photo block
    if form_data.get('photo_url'):
        try:
            img_path = fetch_image_as_jpeg(form_data['photo_url'], target_size=(400, 400))
            photo_elem = RLImage(img_path, width=1.1*inch, height=1.1*inch)
        except:
            photo_elem = Table([["Photo"]], colWidths=[1.1*inch], rowHeights=[1.1*inch], style=[('BOX', (0,0), (-1,-1), 1, colors.black), ('ALIGN', (0,0), (-1,-1), 'CENTER'), ('VALIGN', (0,0), (-1,-1), 'MIDDLE')])
    else:
        photo_elem = Table([["Photo"]], colWidths=[1.1*inch], rowHeights=[1.1*inch], style=[('BOX', (0,0), (-1,-1), 1, colors.black), ('ALIGN', (0,0), (-1,-1), 'CENTER'), ('VALIGN', (0,0), (-1,-1), 'MIDDLE')])
            
    sign_text_str = (
        "<i>Signature of the Official</i><br/><br/>"
        "<b><i>Recommended/Not Recommended</i></b><br/>"
        "<i>(By the Head of the Institute)</i><br/><br/><br/><br/>"
        "<i>(Signature of head of institution)</i><br/>"
        "<i>Name & Designation with Seal</i>"
    )
    
    if form_data.get('sign_url'):
        try:
            img_path = fetch_image_as_jpeg(form_data['sign_url'], target_size=(600, 200))
            sign_img = RLImage(img_path, width=1.5*inch, height=0.5*inch, hAlign='CENTER')
            sign_elem = [sign_img, Paragraph(sign_text_str, center_st)]
        except:
            sign_elem = Paragraph("<br/><br/><br/>" + sign_text_str, center_st)
    else:
        sign_elem = Paragraph("<br/><br/><br/>" + sign_text_str, center_st)
            
    t4_data = [
        [photo_elem, "", sign_elem]
    ]
    
    t4 = Table(t4_data, colWidths=[usable_w*0.3, usable_w*0.25, usable_w*0.45])
    t4.setStyle(TableStyle([
        ('VALIGN', (0,0), (-1,-1), 'TOP'),
        ('ALIGN', (2,0), (2,-1), 'CENTER'),
        ('ALIGN', (0,0), (0,-1), 'LEFT')
    ]))
    story.append(t4)
    
    doc.build(story)
    buf.seek(0)
    return buf

