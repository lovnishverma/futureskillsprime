from docx import Document
from docx.shared import Inches
from PIL import Image
from io import BytesIO

doc = Document()
doc.add_paragraph("Test image below:")

# Create a dummy image
img = Image.new('RGB', (100, 100), color = 'red')
img_buf = BytesIO()
img.save(img_buf, format="JPEG")
img_buf.seek(0)

doc.add_picture(img_buf, width=Inches(1.0))
doc.save("test_image.docx")
print("Saved test_image.docx")
